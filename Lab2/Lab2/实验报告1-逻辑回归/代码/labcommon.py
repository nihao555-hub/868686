# -*- coding: utf-8 -*-
"""通用工具：中文 matplotlib 字体配置 + 《机器学习》实验报告 docx 构建器。

报告模板结构参考仓库中提供的 实验报告N.doc：
封面（班级/学号/姓名/指导教师/学院/日期）+ 报告正文（实验名称、实验地点、
所用工具软件及环境、实验目的、实验内容、实验步骤、程序设计的核心代码、
实验结果、实验体会、教师评价表）。
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

# ----------------------------------------------------------------------------
# 中文字体
# ----------------------------------------------------------------------------
_FONT_CANDIDATES = [
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
]
ZH_FONT_PATH = next((p for p in _FONT_CANDIDATES if os.path.exists(p)), None)
if ZH_FONT_PATH:
    font_manager.fontManager.addfont(ZH_FONT_PATH)
    _name = font_manager.FontProperties(fname=ZH_FONT_PATH).get_name()
    plt.rcParams["font.sans-serif"] = [_name, "WenQuanYi Zen Hei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["figure.dpi"] = 110
plt.rcParams["savefig.dpi"] = 130
plt.rcParams["font.size"] = 11


# ----------------------------------------------------------------------------
# 文件定位
# ----------------------------------------------------------------------------
REPO = "/home/ubuntu/868686"


def find_file(name, root=REPO):
    """在仓库中按文件名查找第一个匹配，返回绝对路径。"""
    for dp, ds, fs in os.walk(root):
        if ".git" in dp:
            continue
        if name in fs:
            return os.path.join(dp, name)
    raise FileNotFoundError(name)


def find_dir(name, root=REPO):
    for dp, ds, fs in os.walk(root):
        if ".git" in dp:
            continue
        if os.path.basename(dp) == name:
            return dp
    raise FileNotFoundError(name)


# ----------------------------------------------------------------------------
# docx 报告构建器
# ----------------------------------------------------------------------------
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_CN_FONT = "宋体"
_CN_FONT_HEI = "黑体"


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


class Report:
    """《机器学习》实验报告。"""

    def __init__(self, exp_no, exp_name, location="信息318",
                 author="________", sid="________________", cls="________________",
                 teacher="戴蕾"):
        self.doc = Document()
        self.exp_no = exp_no
        self.exp_name = exp_name
        self.location = location
        self.author = author
        self.sid = sid
        self.cls = cls
        self.teacher = teacher
        self._setup_styles()
        self._cover()
        self._header()

    # ---- 样式 ----
    def _setup_styles(self):
        st = self.doc.styles["Normal"]
        st.font.name = _CN_FONT
        st.font.size = Pt(12)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), _CN_FONT)
        for section in self.doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)

    def _p(self, text="", size=12, bold=False, align=None, font=_CN_FONT,
           color=None, italic=False, space_after=6):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.25
        return p

    # ---- 封面 ----
    def _cover(self):
        for _ in range(3):
            self._p("")
        self._p("《机器学习》", size=30, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                font=_CN_FONT_HEI)
        self._p("实验报告本", size=30, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                font=_CN_FONT_HEI)
        for _ in range(4):
            self._p("")
        info = [("班    级", self.cls), ("学    号", self.sid),
                ("姓    名", self.author), ("指导教师", self.teacher)]
        tbl = self.doc.add_table(rows=len(info), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (k, v) in enumerate(info):
            c0 = tbl.cell(i, 0)
            c1 = tbl.cell(i, 1)
            self._fill_cell(c0, k, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
            self._fill_cell(c1, v, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
            c0.width = Inches(1.8)
            c1.width = Inches(2.6)
        for _ in range(5):
            self._p("")
        self._p("信息科学与工程学院", size=16, bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER, font=_CN_FONT_HEI)
        self._p("2025 年 6 月", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.doc.add_page_break()

    def _fill_cell(self, cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                   size=12, font=_CN_FONT, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        if color is not None:
            run.font.color.rgb = RGBColor(*color)

    # ---- 正文表头 ----
    def _header(self):
        self._p("《机器学习》实验报告", size=20, bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER, font=_CN_FONT_HEI, space_after=12)
        meta = [
            ("实验名称", f"实验{self.exp_no}：{self.exp_name}"),
            ("实验地点", self.location),
            ("所用工具软件及环境", "Python 3.12；NumPy / Pandas / "
             "scikit-learn / Matplotlib / SciPy"),
        ]
        for k, v in meta:
            p = self.doc.add_paragraph()
            r1 = p.add_run(f"{k}：")
            r1.bold = True
            r1.font.name = _CN_FONT_HEI
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), _CN_FONT_HEI)
            r2 = p.add_run(v)
            r2.font.name = _CN_FONT
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), _CN_FONT)
            for r in (r1, r2):
                r.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(4)

    # ---- 章节 ----
    def h(self, title):
        self._p(title, size=15, bold=True, font=_CN_FONT_HEI, space_after=6)

    def h2(self, title):
        self._p(title, size=13, bold=True, font=_CN_FONT_HEI, space_after=4)

    def para(self, text, **kw):
        return self._p(text, **kw)

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            run = p.add_run(it)
            run.font.size = Pt(12)
            run.font.name = _CN_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), _CN_FONT)
            p.paragraph_format.space_after = Pt(2)

    def code(self, text, caption=None):
        if caption:
            self._p(caption, size=11, italic=True, color=(0x55, 0x55, 0x55),
                    space_after=2)
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        # 浅灰背景
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
        return p

    def image(self, path, width_in=5.6, caption=None):
        self.doc.add_picture(path, width=Inches(width_in))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            self._p(caption, size=10.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                    color=(0x55, 0x55, 0x55), space_after=10)

    def table(self, header, rows, caption=None, col_widths=None,
              first_col_bold=False):
        if caption:
            self._p(caption, size=10.5, italic=True, color=(0x55, 0x55, 0x55),
                    space_after=2)
        t = self.doc.add_table(rows=1, cols=len(header))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for j, htext in enumerate(header):
            self._fill_cell(hdr[j], str(htext), bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, size=11,
                            font=_CN_FONT_HEI)
            _set_cell_bg(hdr[j], "D9E1F2")
        for row in rows:
            cells = t.add_row().cells
            for j, val in enumerate(row):
                self._fill_cell(cells[j], str(val),
                                bold=(first_col_bold and j == 0),
                                align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        if col_widths:
            for j, w in enumerate(col_widths):
                for r in t.rows:
                    r.cells[j].width = Inches(w)
        self._p("", space_after=4)
        return t

    # ---- 教师评价表 ----
    def teacher_eval(self):
        self.h("八、教师评价")
        header = ["评分标准", "分值", "得分"]
        rows = [
            ["整体任务完成度", "20", ""],
            ["实验步骤清晰度", "20", ""],
            ["算法设计合理度", "20", ""],
            ["实验结果正确度", "20", ""],
            ["心得体会深度", "20", ""],
            ["合计", "100", ""],
        ]
        self.table(header, rows, col_widths=[3.0, 1.2, 1.2])
        self._p("教师签名：____________________", size=12,
                align=WD_ALIGN_PARAGRAPH.RIGHT)

    def save(self, path_docx):
        os.makedirs(os.path.dirname(path_docx), exist_ok=True)
        self.doc.save(path_docx)
        return path_docx


def to_pdf(docx_path, outdir=None):
    """用 libreoffice 将 docx 转为 pdf，返回 pdf 路径。"""
    import subprocess
    outdir = outdir or os.path.dirname(docx_path)
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir,
         docx_path],
        check=True, capture_output=True,
    )
    return os.path.splitext(docx_path)[0] + ".pdf"
